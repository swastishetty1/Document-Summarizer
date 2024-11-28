import streamlit as st
import os
from langchain.prompts import PromptTemplate
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains.summarize import load_summarize_chain
from langchain_community.chat_models.oci_generative_ai import ChatOCIGenAI
from pypdf import PdfReader
from io import BytesIO
from typing import Any, Dict, List
import re
from langchain.docstore.document import Document
import docx
 
@st.cache_data
def parse_txt(file: BytesIO) -> str:
    """
    Parses the uploaded .txt file and returns its content as a string.
   
    :param file: The uploaded file (Streamlit file uploader object)
    :return: Content of the .txt file as a string
    """
    text = file.read().decode("utf-8")  # Decode the uploaded file content as a string
    return text
 
@st.cache_data
def parse_pdf(file: BytesIO) -> List[str]:
    pdf = PdfReader(file)
    output = []
    for page in pdf.pages:
        text = page.extract_text()
        # Merge hyphenated words
        text = re.sub(r"(\w+)-\n(\w+)", r"\1\2", text)
        # Fix newlines in the middle of sentences
        text = re.sub(r"(?<!\n\s)\n(?!\s\n)", " ", text.strip())
        # Remove multiple newlines
        text = re.sub(r"\n\s*\n", "\n\n", text)
        output.append(text)
    return output

@st.cache_data
def parse_doc(file: BytesIO) -> List[str]: 
    doc = docx.Document(file)
    output = []
    for paragraph in doc.paragraphs:
        output.append(paragraph.text)
    return [text for text in output if text.strip()]

@st.cache_data
def text_to_docs(text: str,chunk_size,chunk_overlap) -> List[Document]:
    """Converts a string or list of strings to a list of Documents
    with metadata."""
    print(chunk_size)
    print(chunk_overlap)
    if isinstance(text, str):
        # Take a single string as one page
        text = [text]
    page_docs = [Document(page_content=page) for page in text]
 
    # Add page numbers as metadata
    for i, doc in enumerate(page_docs):
        doc.metadata["page"] = i + 1
 
    # Split pages into chunks
    doc_chunks = []
 
    for doc in page_docs:
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""],
            chunk_overlap=chunk_overlap,
        )
        chunks = text_splitter.split_text(doc.page_content)
        for i, chunk in enumerate(chunks):
            doc = Document(
                page_content=chunk, metadata={"page": doc.metadata["page"], "chunk": i}
            )
            # Add sources a metadata
            doc.metadata["source"] = f"{doc.metadata['page']}-{doc.metadata['chunk']}"
            doc_chunks.append(doc)
    return doc_chunks
 
 
def custom_summary(docs, llm, custom_prompt, chain_type, num_summaries):
    print("I am inside custom summary")
    custom_prompt = custom_prompt + """:\n {text}"""
    print("custom Prompt is ------>")
    print(custom_prompt)
    COMBINE_PROMPT = PromptTemplate(template=custom_prompt, input_variables = ["text"])
    print("combine Prompt is ------>")
    print(COMBINE_PROMPT)
    MAP_PROMPT = PromptTemplate(template="Summarize:\n{text}", input_variables=["text"])
    print("MAP_PROMPT Prompt is ------>")
    print(MAP_PROMPT)
    if chain_type == "map_reduce":
        chain = load_summarize_chain(llm,chain_type=chain_type,
                                     map_prompt=MAP_PROMPT,
                                     combine_prompt=COMBINE_PROMPT)
    else:
        chain = load_summarize_chain(llm,chain_type=chain_type)
    print("Chain is --->")
    print(chain)
    summaries = []
    for i in range(num_summaries):
        summary_output = chain({"input_documents": docs}, return_only_outputs=True)["output_text"]
        print("Summaries------------->")
        print(summary_output)
        summaries.append(summary_output)
   
    return summaries
 
 
def main():
    st.set_page_config(layout="wide")
    hide_streamlit_style = """
            <style>
            [data-testid="stToolbar"] {visibility: hidden !important;}
            footer {visibility: hidden !important;}
            </style>
            """
    st.markdown(hide_streamlit_style, unsafe_allow_html=True)
    st.title("Document Summarization App")
         
    user_prompt = st.text_input("Enter the document summary prompt", value= "Compose a brief summary of this text. ")
                                             
    opt = "Upload-own-file"
    pages = None
    if opt == "Upload-own-file":
        uploaded_file = st.file_uploader(
        "**Upload a file :**",
            type=["pdf","txt","docx"],
            )
        if uploaded_file:
            if uploaded_file.name.endswith(".txt"):
                doc = parse_txt(uploaded_file)
            elif uploaded_file.name.endswith(".pdf"):
                doc = parse_pdf(uploaded_file)
            elif uploaded_file.name.endswith(".docx"):
                doc = parse_doc(uploaded_file)
            pages = text_to_docs(doc, 2000, 200)
            print("Pages are here")
            print(pages)
 
 
            page_holder = st.empty()
            if pages:
                print("Inside if Pages")
                st.success("PDF loaded successfully!")
                with page_holder.expander("File Content", expanded=False):
                    pages
 
       
                llm = ChatOCIGenAI(  model_id="meta.llama-3.1-405b-instruct",
    service_endpoint="https://inference.generativeai.us-chicago-1.oci.oraclecloud.com",
    compartment_id="ocid1.tenancy.oc1..aaaaaaaaweuxa6ovnhihlbpolrh3jrdpasnnukjd5x5slxcekzwsdigsayza",
    model_kwargs={"temperature": 0, "max_tokens": 4000}   
)
 
                if st.button("Summarize"):
                    with st.spinner('Summarizing....'):
                        result = custom_summary(pages, llm, user_prompt,"stuff", 1)
                        st.write("Summary:")
                    for summary in result:
                        st.write(summary)
            else:
                st.warning("No file found. Upload a file to summarize!")
           
if __name__=="__main__":
    main()